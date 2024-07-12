# routes.py

import pandas as pd
from flask import Flask, flash, session, render_template, request, jsonify, redirect, send_file, url_for
from feedback_analyzer import FeedbackAnalyzer
from sklearn.model_selection import train_test_split
from io import BytesIO
from docx import Document
from datetime import datetime
from admin import validate_admin, Admin  

def initialize_routes(app, chatbot):

    @app.route('/')
    def index():
        return render_template('index.html')
    
    @app.route('/login', methods=['GET', 'POST'])
    def login():
        if request.method == 'POST':
            username = request.form['username']
            password = request.form['password']
            
            if validate_admin(username, password):
                session['admin'] = username
                return redirect(url_for('feedback_report'))  # Replace with your desired redirect
            else:
                flash('Invalid username or password', 'error')

        return render_template('login.html')
    @app.route('/logout')
    def logout():
        session.pop('admin', None)  # Remove 'admin' from session
        return redirect(url_for('login'))

    @app.route('/clear_session')
    def clear_session():
        session.clear()  # Clear all session variables
        return redirect(url_for('login'))

    @app.route('/change_password', methods=['GET', 'POST'])
    def change_password():
        if request.method == 'POST':
            username = request.form['username']
            current_password = request.form['currentPassword']
            new_password = request.form['newPassword']
            confirm_password = request.form['confirmPassword']

            if new_password != confirm_password:
                flash('New passwords do not match', 'error')
            elif not validate_admin(username, current_password):
                flash('Current password is incorrect', 'error')
            else:
                # Process password change logic here
                # Example: update password in database or any other action
                flash('Password changed successfully', 'success')
                return redirect(url_for('login'))

        return render_template('change_pass.html')
    @app.route('/chat_1')
    def chat_1():
        return render_template('chat_1.html')

    @app.route('/chat_2')
    def chat_2():
        return render_template('chat_2.html')

    @app.route('/chat_3')
    def chat_3():
        return render_template('chat_3.html')

    @app.route('/rate', methods=['GET', 'POST'])
    def rate_page():
        if request.method == 'POST':
            chatbot.end_chat()  # Trigger end_chat when user submits rating
            return redirect('/')
        else:
            return render_template('rate.html')

    @app.route('/chat', methods=['POST'])
    def chat():
        try:
            prompt = request.form['prompt']
            sentiment = request.form.get('sentiment', 'neutral')  # Default to neutral if not provided
            
            print(f"Received prompt: {prompt}, sentiment: {sentiment}")  # Log received data
            
            if prompt.upper() == 'END CHAT':
                chatbot.end_chat()  # Trigger end_chat when user explicitly ends chat
                return 'END CHAT'
            elif prompt.upper() == 'RESTAURANT NAME':
                response = chatbot.get_restaurant_name()
            elif prompt.upper() == 'RESTAURANT PROFILE':
                response = chatbot.get_restaurant_profile()
            elif prompt.upper() == 'MENU':
                response = chatbot.get_menu()
            else:
                response = chatbot.generate_response(prompt)
            
            # Generate suggestions based on sentiment
            suggestions = chatbot.generate_suggestion_prompt(prompt, sentiment)
            
            print(f"Response: {response}")
            print(f"Suggestions: {suggestions}")
            
            return jsonify({"response": response, "suggestions": suggestions})
        except Exception as e:
            print(f"Error in chat: {e}")
            return jsonify({"error": str(e)})

    @app.route('/evaluate_model', methods=['GET'])
    def evaluate_model():
        if chatbot is None:
            return jsonify({"error": "Chatbot initialization failed"})
        try:
            hyperparameters = {'alpha': [0.1, 0.5, 1.0, 1.5, 2.0]}
            df = pd.read_csv('annotated_empatheticdialogues.csv')
            combined_text = df['context'].astype(str) + ' ' + df['prompt'].astype(str) + ' ' + df['utterance'].astype(str) + ' ' + df['emotion'].astype(str)
            labels = df['emotion'].tolist()
            X_train, X_test, y_train, y_test = train_test_split(combined_text, labels, test_size=0.2, random_state=42)
            chatbot.model.train_model(file_path='annotated_empatheticdialogues.csv')  # Call train_model without additional arguments
            results = chatbot.model.evaluate_model(X_test, y_test)
            return jsonify({"results": results})
        except Exception as e:
            print(f"Error evaluating model: {e}")
            return jsonify({"error": str(e)})

    @app.route('/analyze_feedback', methods=['GET'])
    def analyze_feedback():
        if chatbot is None:
            return jsonify({"error": "Chatbot initialization failed"})
        try:
            feedback_analyzer = FeedbackAnalyzer(chatbot.feedback)
            categorized_feedback, summary = feedback_analyzer.interpret_feedback()
            return render_template('report.html', feedback_data=categorized_feedback, summary=summary)
        except Exception as e:
            print(f"Error analyzing feedback: {e}")
            return jsonify({"error": str(e)})

    @app.route('/feedback_report')
    def feedback_report():
        if 'admin' not in session:
            return redirect(url_for('login'))
        if chatbot is None:
            return jsonify({"error": "Chatbot initialization failed"})
        else:
            feedback_analyzer = FeedbackAnalyzer(chatbot.feedback)
            categorized_feedback, summary = feedback_analyzer.interpret_feedback()
            return render_template('report.html', feedback_data=categorized_feedback, summary=summary)

    @app.route('/download_report')
    def download_report():
        if 'admin' not in session:
            return redirect(url_for('login'))
        if chatbot is None:
            return jsonify({"error": "Chatbot initialization failed"})
        try:
            feedback_analyzer = FeedbackAnalyzer(chatbot.feedback)
            categorized_feedback, summary = feedback_analyzer.interpret_feedback()
            
            def create_docx_report(categorized_feedback, summary):
                document = Document()
                document.add_heading('Feedback Report', 0)
                
                # Add date and timestamp
                current_datetime = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                document.add_paragraph(f"Generated on: {current_datetime}")
                
                for category, feedbacks in categorized_feedback.items():
                    category_title = category.replace('_', ' ').title()
                    document.add_heading(f"{category_title} Feedback", level=1)
                    
                    if category != 'suggestions':
                        document.add_heading('Compliments', level=2)
                        for feedback in feedbacks['compliments']:
                            document.add_paragraph(f"- {feedback['message']}")
                        
                        document.add_heading('Complaints', level=2)
                        for feedback in feedbacks['complaints']:
                            document.add_paragraph(f"- {feedback['message']}")
                    else:
                        document.add_heading('Suggestions', level=2)
                        for feedback in feedbacks:
                            document.add_paragraph(f"- {feedback['message']}")
                    
                    document.add_paragraph('')
                
                document.add_heading('Summary', level=1)
                document.add_paragraph(summary)
                
                return document
            
            docx_report = create_docx_report(categorized_feedback, summary)
            
            buffer = BytesIO()
            docx_report.save(buffer)
            buffer.seek(0)
            return send_file(buffer, as_attachment=True, download_name='report.docx', mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        except Exception as e:
            print(f"Error downloading report: {e}")
            return jsonify({"error": str(e)})
